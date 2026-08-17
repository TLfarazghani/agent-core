"""create_docx handler — Windows backend via python-docx.

Schema (research doc §3.4): ``{ "title": str, "sections": [{heading, body}] }``.
Writes the document to the output dir and returns the created path as the
tool result.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any, Callable

from docx import Document

from ._paths import default_output_dir, slugify, unique_path

TOOL_NAME = "create_docx"


def _as_body_text(value: Any) -> str:
    """Accept a string, a list of lines, or None; always return a string."""
    if value is None:
        return ""
    if isinstance(value, list):
        return "\n".join(str(item) for item in value)
    return str(value)


def _normalize_sections(sections: Any) -> list[dict[str, Any]]:
    """Accept a single section object or a list of them; fill missing fields."""
    if isinstance(sections, dict):
        sections = [sections]
    return [
        {"heading": str(section.get("heading", "")), "body": _as_body_text(section.get("body"))}
        for section in sections
    ]


def make_handler(output_dir: Path | None = None) -> Callable[[dict[str, Any]], str]:
    out_dir = output_dir or default_output_dir()

    def handler(arguments: dict[str, Any]) -> str:
        title = arguments["title"]
        sections = _normalize_sections(arguments["sections"])

        doc = Document()
        doc.add_heading(title, level=0)
        for section in sections:
            doc.add_heading(section["heading"], level=1)
            doc.add_paragraph(section["body"])

        out_dir.mkdir(parents=True, exist_ok=True)
        path = unique_path(out_dir, slugify(title), "docx")
        doc.save(str(path))
        return str(path)

    return handler